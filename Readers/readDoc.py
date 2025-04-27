#!pip install python-docx 
# does not read header/footer normally
# does not read images
# reads tables and paras

from docx import Document

def get_data_from_text(doc):
	a=[]
	for para in doc.paragraphs:
		p=para.text
		if len(p)!=0:
			a.append(p)
	k=[]
	for i in a:
		for m in i.split('\n'):
			if len(m.strip())!=0:
				k.append(m.strip())
	return k


def get_data_from_table(doc):
	a=[]
	for table in doc.tables:
		for row in table.rows:
			for cell in row.cells:
				s = cell.text
				if len(s) != 0:
					a.append(cell.text)
	k=[]
	for i in a:
		for m in i.split('\n'):
			if len(m.strip())!=0:
				k.append(m.strip())
	return k
def total_data(doc):
	return get_data_from_table(doc) + get_data_from_text(doc)

def readDoc(filename):
	if not(filename.endswith('.docx')):
		filename=filename+'.docx'
	try:
		doc = Document(filename)
		return " ".join(total_data(doc))
	except Exception as e:
		print('File location missing or refer error '+str(e))
		return ""
# print(readDoc('w'))
